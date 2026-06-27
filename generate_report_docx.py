from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE_DIR = Path(r"d:/大数据分析实践报告")
OUTPUT_DOCX = BASE_DIR / "智能家居销量数据分析系统实验报告.docx"
ASSET_DIR = BASE_DIR / "assests"

REPORT_DATA = {
    "system_name": "智能家居销量数据分析系统",
    "backend": "Spring Boot 2.7.18、MyBatis Plus、MySQL",
    "frontend": "Vue 3、Element Plus、ECharts、Axios",
    "sales_records": 1415,
    "merged_records": 9999,
    "prediction_records": 1530,
    "feature_records": 11,
    "sales_start": "2022-01-02",
    "sales_end": "2025-11-18",
    "total_sales": "194514852.00",
    "total_orders": "9993",
    "avg_order_amount": "19455.37",
}


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.175)
    section.right_margin = Cm(3.175)

    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    for style_name, font_name, size in [
        ("Heading 1", "黑体", 16),
        ("Heading 2", "黑体", 14),
        ("Heading 3", "黑体", 12),
        ("Heading 4", "黑体", 12),
    ]:
        styles[style_name].font.name = font_name
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        styles[style_name].font.size = Pt(size)


def add_center_paragraph(document: Document, text: str, size: int, bold: bool = False, space_before: float = 0, space_after: float = 0):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    return paragraph


def add_left_paragraph(document: Document, text: str, size: int, space_after: float = 6):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    return paragraph


def add_heading(document: Document, text: str, level: int):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_body(document: Document, text: str):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    return paragraph


def add_caption(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)
    return paragraph


def add_code_block(document: Document, code: str):
    for line in code.strip("\n").split("\n"):
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(10.5)


def add_image(document: Document, filename: str, width_cm: float = 15.5):
    document.add_picture(str(ASSET_DIR / filename), width=Cm(width_cm))


def add_cover_page(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()
    add_center_paragraph(document, "大数据分析课程实践", 22, True, 0, 8)
    add_center_paragraph(document, "课程设计报告", 22, True, 0, 40)
    add_left_paragraph(document, "题    目：智能家居销量数据分析系统", 14)
    add_left_paragraph(document, "学生姓名：刘培冉", 14)
    add_left_paragraph(document, "学    号：23125082009", 14)
    add_left_paragraph(document, "学    院：信息工程学院", 14)
    add_left_paragraph(document, "专    业：数据科学与大数据技术", 14)
    add_left_paragraph(document, "年    级：2023级02班", 14)
    add_left_paragraph(document, "指导教师：胡贞华", 14)
    document.add_page_break()
    add_center_paragraph(document, "2024—2025学年 第 一 学期", 14, False, 0, 12)


def add_section_purpose_and_content(document: Document) -> None:
    add_heading(document, "实践目的", 1)
    add_heading(document, "1.1 时代要求", 2)
    add_body(document, "随着智能家居设备的普及，线上销售平台、线下渠道活动以及用户消费行为不断积累形成结构化经营数据。围绕销售额、订单量、地区分布、产品品类、促销方式和用户画像开展分析，已经成为数据科学课程实践中的典型应用场景。本项目以智能家居销量数据分析系统为载体，构建了从数据清洗、结果存储到前后端展示的完整链路，体现了数据驱动决策在真实业务中的价值。")
    add_body(document, "传统静态表格难以满足多维度经营数据的对比观察需求，而本项目借助 Vue 3 与 ECharts 对销售趋势、地域分布、产品结构、用户画像和随机森林预测结果进行图形化展示，使用户能够在同一系统中快速理解销售变化规律并辅助管理决策。")
    add_body(document, "此外，项目在分析模块中引入随机森林回归模型，对销售结果进行预测并给出特征重要性排序，使课程实践不仅停留在数据可视化层面，还进一步延伸到建模分析与结果解释层面，体现了数据科学与大数据技术专业中“采集—处理—分析—可视化—应用”的综合训练目标。")
    add_heading(document, "1.2 教育要求", 2)
    add_body(document, "通过本项目实践，可以系统训练学生完成一个数据分析系统的设计与实现。学生不仅需要理解前后端分离架构和数据库建模方式，还需要掌握数据清洗、分析结果组织、接口设计、页面可视化以及预测模型结果展示等关键流程。")
    add_body(document, "在技术能力方面，本项目覆盖了 Python 数据处理脚本、MySQL 数据存储、Spring Boot 后端服务、Vue 前端开发与 ECharts 图表展示等内容；在工程能力方面，本项目要求学生关注模块划分、接口协作、数据一致性与系统运行效果，从而实现从理论课程到工程落地的有效衔接。")

    add_heading(document, "实践内容", 1)
    add_heading(document, "2.1 使用技术栈", 2)
    add_body(document, f"本项目后端采用 {REPORT_DATA['backend']} 作为核心支撑技术，通过 MyBatis Plus 完成数据表访问与业务接口实现，并利用 MySQL 数据库存储用户信息、销售汇总、地域分析、产品分析、促销分析、用户画像、预测结果和特征重要性等业务数据。")
    add_body(document, f"前端部分采用 {REPORT_DATA['frontend']} 构建可视化系统，支持登录注册、首页展示、图表分析和数据管理功能。前端通过 Axios 与后端交互，使可视化界面与数据库中的分析结果保持一致。")
    add_body(document, "项目还配套提供 Python 数据分析流程，按照“数据清洗—从 MySQL 读取分析—结果写回 MySQL”的方式组织处理逻辑，为系统中的可视化与预测页面提供稳定、可持续的数据来源。")
    add_heading(document, "2.2 数据收集", 2)
    add_body(document, f"本系统的数据基础来源于项目内置的数据处理流程和数据库初始化结果。根据当前数据库实际情况，清洗汇总后的 merged_data 表共包含 {REPORT_DATA['merged_records']} 条记录，销售汇总表 sales_summary 共 {REPORT_DATA['sales_records']} 条记录，数据时间范围覆盖 {REPORT_DATA['sales_start']} 至 {REPORT_DATA['sales_end']}，能够较完整地反映智能家居销售活动在时间、地区、品类与促销维度上的变化情况。")
    add_heading(document, "2.3 系统功能模块", 2)
    add_body(document, "系统功能主要包括用户认证、数据管理和结果可视化三部分。用户认证模块支持登录、注册和路由守卫；数据管理模块支持对多个分析结果表进行增删改查；可视化模块则围绕销售趋势、地域分析、产品品类、用户画像、销售渠道以及随机森林预测等主题进行展示。")


def add_section_process(document: Document) -> None:
    add_heading(document, "实践过程", 1)
    add_heading(document, "3.1 准备阶段", 2)
    add_body(document, "本项目首先完成 MySQL 数据库环境准备，并导入系统自带的 smart_home.sql 数据脚本。当前数据库中已验证存在销售汇总、地域分析、产品分析、促销分析、用户画像、预测结果、特征重要性等多张业务表，为后续后端接口与前端页面提供稳定数据支撑。")
    add_body(document, "后端服务通过 Spring Boot 打包结果启动，前端通过 Vue 开发服务器运行，系统最终以前后端分离方式完成调试。经实际验证，后端接口能够正常返回销售趋势数据，前端页面也可在浏览器中成功加载。")
    add_heading(document, "3.2 数据处理", 2)
    add_body(document, "项目中的 Python 分析流程将数据处理分为三步：首先对原始数据进行清洗与合并，然后从 MySQL 中读取清洗后的数据执行分析，最后将分析结果统一写回数据库。该流程使系统的可视化结果与数据库内容形成闭环，降低了人工整理分析表的重复劳动。")
    add_heading(document, "3.3 主要页面展示", 2)
    add_body(document, "系统首页提供登录入口，用户可通过输入账号密码进入系统；当用户尚未拥有账号时，也可以通过注册页面完成基本信息填写并创建账户。")
    add_image(document, "微信图片_20260616153634_159_330.png")
    add_caption(document, "图3-3-1：登录界面")
    add_body(document, "注册界面与登录流程配套，为系统提供基础的用户创建能力。通过该模块可以验证前端表单与后端认证接口的联动效果。")
    add_image(document, "微信图片_20260616153652_160_330.png")
    add_caption(document, "图3-3-2：注册界面")
    add_body(document, "进入系统后，用户可以在首页或数据大屏中整体查看智能家居销量分析结果，系统将多个维度的经营数据集中呈现，形成统一的可视化入口。")
    add_image(document, "微信图片_20260616153701_161_330.png")
    add_caption(document, "图3-3-3：系统首页界面")
    add_heading(document, "3.4 可视化分析模块展示", 2)
    add_heading(document, "3.4.1 销售趋势与地域分析", 3)
    add_body(document, "销售趋势分析模块围绕销售额随时间变化的规律展开，能够展示阶段性波动情况；地域分析模块则反映不同地区之间的销售规模差异。在当前数据库中，地域分析结果共 9 条，销售额较高的地区能够直观反映智能家居产品在区域市场中的表现差别。")
    add_image(document, "微信图片_20260616153711_162_330.png")
    add_caption(document, "图3-4-1：销售趋势或地域分析界面")
    add_heading(document, "3.4.2 产品、促销与用户画像分析", 3)
    add_body(document, "产品分析模块主要统计品类、品牌、销售量与平均价格等指标；促销分析模块用于比较不同促销方式的平均折扣率与销售贡献；用户画像分析则聚焦性别、会员类型、用户数量与平均消费水平等信息。项目当前数据库中 product_analysis 表包含 15 条记录，promotion_analysis 表包含 7 条记录，user_profile 表包含 4 条记录，说明系统已具备多角度分析能力。")
    add_image(document, "微信图片_20260616153718_163_330.png")
    add_caption(document, "图3-4-2：产品或用户画像分析界面")
    add_heading(document, "3.4.3 数据管理与预测展示", 3)
    add_body(document, "除展示结果外，系统还提供数据管理模块用于维护分析结果表内容，支持查看、新增、修改和删除操作。预测模块采用随机森林回归模型展示销售预测结果，并给出特征重要性。当前 prediction_result 表已有 1530 条记录，feature_importance 表已有 11 条记录，为分析结论提供了建模依据。")
    add_image(document, "微信图片_20260616153727_164_330.png")
    add_caption(document, "图3-4-3：数据管理或预测分析界面")


def add_section_code_and_summary(document: Document) -> None:
    add_heading(document, "3.7 部分代码展示", 2)
    add_heading(document, "3.7.1 后端数据源配置代码展示", 3)
    add_code_block(document, '''server:
  port: 8080
spring:
  datasource:
    driver-class-name: com.mysql.cj.jdbc.Driver
    url: jdbc:mysql://127.0.0.1:3306/smart_home?useUnicode=true&characterEncoding=utf8&useSSL=false&serverTimezone=Asia/Shanghai&allowPublicKeyRetrieval=true
    username: root''')
    add_heading(document, "3.7.2 前端路由与页面组织代码展示", 3)
    add_code_block(document, '''const routes = [
  { path: '/login', component: Login },
  { path: '/dashboard', component: Dashboard },
  { path: '/visualization/sales-trend', component: SalesTrend },
  { path: '/visualization/region-analysis', component: RegionAnalysis },
  { path: '/data-management', component: DataManagement }
]''')
    add_heading(document, "3.7.3 随机森林结果展示说明", 3)
    add_body(document, "预测模块基于随机森林回归模型组织结果展示，系统不仅保存预测值与实际值，还输出误差与特征重要性排序，使模型分析结果能够在前端页面中被直接查看与解释。")
    add_heading(document, "实践体会", 1)
    add_body(document, "通过本次智能家居销量数据分析系统实验，我对一个完整数据分析项目的实现过程有了更清晰的认识。与单纯完成算法题或静态图表不同，该项目要求将数据库、后端接口、前端页面与分析流程联结起来，使我体会到工程实现与数据分析在真实应用中的紧密关系。")
    add_body(document, "在阅读并运行项目的过程中，我进一步理解了前后端分离系统中“数据表—接口—页面”之间的映射关系，也认识到数据分析结果只有在结构清晰、展示直观的情况下，才能真正服务于业务判断。尤其是随机森林预测模块的加入，使我对“分析”与“建模解释”之间的差异有了更加具体的认识。")
    add_body(document, "总体来看，本次课程实践不仅提高了我对 Spring Boot、Vue 和 MySQL 等技术组合的理解，也锻炼了我从系统整体出发分析问题、组织结果与表达结论的能力，为后续更复杂的数据系统设计打下了基础。")


def main() -> None:
    document = Document()
    set_doc_defaults(document)
    add_cover_page(document)
    add_section_purpose_and_content(document)
    add_section_process(document)
    add_section_code_and_summary(document)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == '__main__':
    main()
