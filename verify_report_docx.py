from pathlib import Path
from docx import Document

REPORT_PATH = Path(r"d:/大数据分析实践报告/智能家居销量数据分析系统实验报告.docx")


def main() -> None:
    doc = Document(str(REPORT_PATH))
    text = "\n".join(p.text for p in doc.paragraphs)
    checks = [
        "智能家居销量数据分析系统",
        "实践目的",
        "实践内容",
        "实践过程",
        "实践体会",
        "图3-3-1：登录界面",
        "prediction_result 表已有 1530 条记录",
    ]
    for item in checks:
        print(item, item in text)
    print("paragraphs", len(doc.paragraphs))
    rels = [r for r in doc.part._rels.values() if "image" in r.reltype]
    print("images", len(rels))


if __name__ == "__main__":
    main()
